import os
import math
import threading
import time
from tkinter import Tk, Button, Label, filedialog, messagebox, StringVar
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

# WordprocessingML 命名空间
W_NAMESPACE = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# 1. 处理 Word 文档的核心功能
def load_docx(file_path):
    return Document(file_path)

def format_table(table):
    for row in table.rows:
        max_lines = 1
        for cell in row.cells:
            cell_text = cell.text.strip()
            lines = estimate_line_count(cell_text, max_chars_per_line=20)
            max_lines = max(max_lines, lines)

            for paragraph in cell.paragraphs:
                if paragraph.runs:
                    run = paragraph.runs[0]
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        set_dynamic_row_height(row, max_lines)

def estimate_line_count(text, max_chars_per_line=20):
    if not text:
        return 1
    lines = text.count('\n') + 1
    estimated_lines = math.ceil(len(text) / max_chars_per_line)
    return max(lines, estimated_lines)

def set_dynamic_row_height(row, num_lines):
    base_height = 300
    total_height = base_height * num_lines

    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(f'{{{W_NAMESPACE}}}val', str(total_height))
    trHeight.set(f'{{{W_NAMESPACE}}}hRule', 'auto')
    trPr.append(trHeight)

def remove_empty_rows(table):
    for row in table.rows:
        if all(cell.text.strip() == '' for cell in row.cells):
            tbl = table._tbl
            tbl.remove(row._tr)

def get_revised_filename(file_path):
    base, ext = os.path.splitext(file_path)
    return f"{base}_revised{ext}"

def process_docx(input_path):
    if not os.path.isfile(input_path):
        raise FileNotFoundError("指定的文件不存在。请检查路径是否正确。")

    doc = load_docx(input_path)
    for table in doc.tables:
        format_table(table)
        remove_empty_rows(table)

    output_path = get_revised_filename(input_path)
    doc.save(output_path)
    return output_path

# 2. 创建动画效果
def animate(label_var, stop_event):
    animation = ['|', '/', '-', '\\']  # 旋转沙漏动画
    idx = 0
    while not stop_event.is_set():
        label_var.set(f"正在处理中... {animation[idx % len(animation)]}")
        idx += 1
        time.sleep(0.2)
    label_var.set("处理完成！")

# 3. 处理文件的线程函数
def process_file_with_animation(file_path, label_var):
    stop_event = threading.Event()
    animation_thread = threading.Thread(target=animate, args=(label_var, stop_event))
    animation_thread.start()

    try:
        output_file = process_docx(file_path)
        stop_event.set()
        messagebox.showinfo("成功", f"格式调整完成！\n文件已保存至：\n{output_file}")
    except Exception as e:
        stop_event.set()
        messagebox.showerror("错误", f"处理文档时出错：{e}")

# 4. 文件选择与处理逻辑
def select_file(label_var):
    file_path = filedialog.askopenfilename(
        title="选择 Word 文件",
        filetypes=[("Word 文档", "*.docx")]
    )
    if file_path:
        threading.Thread(target=process_file_with_animation, args=(file_path, label_var)).start()

# 5. 创建 GUI 界面
def create_gui():
    window = Tk()
    window.title("Word 表格修复工具")
    window.geometry("400x200")

    label = Label(window, text="请选择要处理的 Word 文件", font=("Arial", 12))
    label.pack(pady=10)

    select_button = Button(window, text="选择文件", command=lambda: select_file(status_var), width=20, height=2)
    select_button.pack(pady=5)

    # 状态提示标签（用于显示动画）
    global status_var
    status_var = StringVar()
    status_var.set("")
    status_label = Label(window, textvariable=status_var, font=("Arial", 10), fg="blue")
    status_label.pack(pady=10)

    window.mainloop()

# 6. 启动 GUI
if __name__ == "__main__":
    create_gui()

