import os
import math
import threading
import time
from tkinter import Tk, Button, Label, filedialog, messagebox, StringVar
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

# WordprocessingML namespace
W_NAMESPACE = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# 1. Core functionality for processing Word documents
def load_docx(file_path):
    return Document(file_path)

def format_table(table):
    for row in table.rows:
        max_lines = 1
        for cell in row.cells:
            cell_text = cell.text.strip()
            lines = estimate_line_count(cell_text, max_chars_per_line=20)
            max_lines = max(max_lines, lines)

            for paragraph in cell.paragraphs:
                if paragraph.runs:
                    run = paragraph.runs[0]
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        set_dynamic_row_height(row, max_lines)

def estimate_line_count(text, max_chars_per_line=20):
    if not text:
        return 1
    lines = text.count('\n') + 1
    estimated_lines = math.ceil(len(text) / max_chars_per_line)
    return max(lines, estimated_lines)

def set_dynamic_row_height(row, num_lines):
    base_height = 300
    total_height = base_height * num_lines

    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(f'{{{W_NAMESPACE}}}val', str(total_height))
    trHeight.set(f'{{{W_NAMESPACE}}}hRule', 'auto')
    trPr.append(trHeight)

def remove_empty_rows(table):
    for row in table.rows:
        if all(cell.text.strip() == '' for cell in row.cells):
            tbl = table._tbl
            tbl.remove(row._tr)

def get_revised_filename(file_path):
    base, ext = os.path.splitext(file_path)
    return f"{base}_revised{ext}"

def process_docx(input_path):
    if not os.path.isfile(input_path):
        raise FileNotFoundError("The specified file does not exist. Please check the path.")

    doc = load_docx(input_path)
    for table in doc.tables:
        format_table(table)
        remove_empty_rows(table)

    output_path = get_revised_filename(input_path)
    doc.save(output_path)
    return output_path

# 2. Create animation effect
def animate(label_var, stop_event):
    animation = ['|', '/', '-', '\\']  # Spinning hourglass animation
    idx = 0
    while not stop_event.is_set():
        label_var.set(f"Processing... {animation[idx % len(animation)]}")
        idx += 1
        time.sleep(0.2)
    label_var.set("Processing complete!")

# 3. Thread function for file processing
def process_file_with_animation(file_path, label_var):
    stop_event = threading.Event()
    animation_thread = threading.Thread(target=animate, args=(label_var, stop_event))
    animation_thread.start()

    try:
        output_file = process_docx(file_path)
        stop_event.set()
        messagebox.showinfo("Success", f"Formatting complete!\nFile saved to:\n{output_file}")
    except Exception as e:
        stop_event.set()
        messagebox.showerror("Error", f"An error occurred while processing the document: {e}")

# 4. File selection and processing logic
def select_file(label_var):
    file_path = filedialog.askopenfilename(
        title="Select Word Document",
        filetypes=[("Word Documents", "*.docx")]
    )
    if file_path:
        threading.Thread(target=process_file_with_animation, args=(file_path, label_var)).start()

# 5. Create GUI interface
def create_gui():
    window = Tk()
    window.title("Word Table Formatter")
    window.geometry("400x200")

    label = Label(window, text="Please select a Word document to process", font=("Arial", 12))
    label.pack(pady=10)

    select_button = Button(window, text="Select File", command=lambda: select_file(status_var), width=20, height=2)
    select_button.pack(pady=5)

    # Status label for displaying animation
    global status_var
    status_var = StringVar()
    status_var.set("")
    status_label = Label(window, textvariable=status_var, font=("Arial", 10), fg="blue")
    status_label.pack(pady=10)

    window.mainloop()

# 6. Launch GUI
if __name__ == "__main__":
    create_gui()

